CATAGORIES_EXTERNAL_ID = [""]
TEAMS_TXT_FILE = "teams.txt"
USERS_TXT_FILE = "users.txt"
PASSWORDS_DOCX = "passwords.docx" # 密碼紙
iterater_team_id = 0 # 第一支隊伍 ID
else_team_count = 0 # 用 team_id 當計分板名稱的數量

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.table import _Cell
import json
from generator import generatorPassword

def read_file(filename):
    try:
        with open(filename, "r", encoding='utf-8') as file:
            return [line.strip() for line in file]
    except:
        return []

def set_font(cell: _Cell) -> None :
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"  # 設置英文字體
    cell.paragraphs[0].runs[0].font.size = Pt(14)  # 字體大小
    cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微軟正黑體')  # 設置中文字體

def create_word_document(team_names, passwords):
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    # 設置列寬
    for i, width in enumerate([5, 10, 2, 3]):
        table.columns[i].width = Cm(width)
    
    # 添加表頭
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(['簽名', '隊伍名', '帳號', '密碼']):
        hdr_cells[i].text = header_text
        set_font(hdr_cells[i])
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # 添加數據
    for i in range(len(team_names)):
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(['', team_names[i], passwords[i]['username'], passwords[i]['password']]):
            row_cells[i].text = cell_text
            set_font(row_cells[i])
    
    # 設置表格格式
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 設置垂直居中
            cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell.paragraphs[0].paragraph_format.line_spacing = 1.5  # 設置行距
    
    doc.save(PASSWORDS_DOCX)

def create_account_data(team_name, user_name = ''):
    if user_name == '':
        user_name = team_name

    global iterater_team_id
    account = {
        "id": "team{:03}".format(iterater_team_id), # external_id
        "username": "team{:03}".format(iterater_team_id), # 帳號
        "password": generatorPassword(), # 密碼
        "type": "team", # 固定
        "name": user_name, # 後臺名字
        "team_id": "team{:03}".format(iterater_team_id) # 所屬的 team_external_id
    }
    iterater_team_id += 1
    return account

def main():
    team_names = read_file(TEAMS_TXT_FILE)
    user_names = read_file(USERS_TXT_FILE)

    if len(user_names) == 0:
        user_names = [''] * len(team_names)

    accounts = []

    for team_name, user_name in zip(team_names, user_names):
        accounts.append(create_account_data(team_name, user_name))

    for i in range(else_team_count):
        accounts.append(create_account_data("team{:03}".format(iterater_team_id)))

    with open("accounts.json", 'w', encoding='utf-8') as accountsFile:
        json.dump(accounts, accountsFile, indent=2, ensure_ascii=False)

    create_word_document(team_names, accounts)
    print("accounts.json and accounts.docx created successfully")
    
if __name__ == '__main__':
    main()