CATAGORIES_EXTERNAL_ID = ["Second_Contest"]
TEAMS_TXT_FILE = "teams.txt"
USERS_TXT_FILE = "users.txt"
PASSWORDS_DOCX = "passwords.docx" # 密碼紙
iterater_team_id = 300 # 第一支隊伍 ID
else_team_count = 5 # 用 team_id 當計分板名稱的數量

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.oxml import OxmlElement
import json
from generator import generatorPassword

def read_file(filename: str) -> list[str]:
    try:
        with open(filename, "r", encoding='utf-8') as file:
            return [line.strip() for line in file]
    except:
        return []

def set_font(cell: _Cell) -> None:
    cell.paragraphs[0].runs[0].font.name = "Times New Roman"  # 設置英文字體
    cell.paragraphs[0].runs[0].font.size = Pt(14)  # 字體大小
    cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')  # 設置中文字體

def set_cell_border(cell: _Cell, **kwargs):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for edge in ['top', 'left', 'bottom', 'right']:
        if edge in kwargs:
            # 建立邊框元素，使用正確的命名空間
            edge_el = OxmlElement(f'w:{edge}')
            edge_el.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            # edge_el.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))  # 預設邊框粗細
            # edge_el.set(qn('w:space'), str(kwargs[edge].get('space', 0)))
            # edge_el.set(qn('w:color'), kwargs[edge].get('color', 'auto'))
            tcPr.append(edge_el)

def create_word_document(team_names: list[str], accounts: dict) -> None:
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(0.5)  # 左邊界 0.5 公分
    section.right_margin = Cm(0.5)  # 右邊界 0.5 公分
    section.top_margin = Cm(0.5)  # 上邊界 0.5 公分
    section.bottom_margin = Cm(0.5)  # 下邊界 0.5 公分

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'

    # 設置列寬
    for i, width in enumerate([5, 5, 5, 2, 2]):
        table.columns[i].width = Cm(width)
    
    # 添加表頭
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(['簽名', '隊伍名', '成員', '帳號', '密碼']):
        hdr_cells[i].text = header_text
        set_font(hdr_cells[i])
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # 若為「帳號」或「密碼」，設定下側為虛線
        if header_text in ['帳號', '密碼']:
            set_cell_border(hdr_cells[i], bottom={'val': 'dashed'})
    
    # 添加數據
    for i in range(len(team_names)):
        row_cells = table.add_row().cells
        for j, cell_text in enumerate(['',
                                       team_names[i],
                                       accounts[i]['name'],
                                       accounts[i]['username'],
                                       accounts[i]['password']]):
            row_cells[j].text = cell_text
            set_font(row_cells[j])

            # 為帳號和密碼欄位設置虛線邊框
            if j == 2:
                set_cell_border(row_cells[j], right={'val': 'dashed'})
            if j in [3, 4]:  # 第 4 和第 5 列
                if i == len(team_names) - 1:
                    set_cell_border(row_cells[j],
                                    top={'val': 'dashed'},
                                    left={'val': 'dashed'})
                else:
                    set_cell_border(row_cells[j],
                                    top={'val': 'dashed'},
                                    left={'val': 'dashed'},
                                    bottom={'val': 'dashed'})
    
    # 設置表格格式
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # 設置垂直居中
            cell.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.LEFT
            cell.paragraphs[0].paragraph_format.line_spacing = 1.5  # 設置行距
    
    doc.save(PASSWORDS_DOCX)

def create_account_data(team_name, user_name = None) -> dict:
    if user_name is None:
        user_name = team_name

    global iterater_team_id
    account = {
        "id": "team{:03}".format(iterater_team_id), # external_id
        "username": "team{:03}".format(iterater_team_id), # 帳號
        "password": generatorPassword(), # 密碼
        "type": "team", # 固定
        "name": str(user_name), # 後臺名字
        "team_id": "team{:03}".format(iterater_team_id) # 所屬的 team_external_id
    }
    iterater_team_id += 1
    return account

def main():
    team_names = read_file(TEAMS_TXT_FILE)
    user_names = read_file(USERS_TXT_FILE)

    if len(user_names) == 0:
        user_names = [''] * len(team_names)

    accounts = []

    for team_name, user_name in zip(team_names, user_names):
        accounts.append(create_account_data(team_name, user_name))

    for i in range(else_team_count):
        team_names.append("team{:03}".format(iterater_team_id))
        accounts.append(create_account_data("team{:03}".format(iterater_team_id)))

    with open("accounts.json", 'w', encoding='utf-8') as accountsFile:
        json.dump(accounts, accountsFile, indent=2, ensure_ascii=False)

    create_word_document(team_names, accounts)
    print("accounts.json and accounts.docx created successfully")
    
if __name__ == '__main__':
    main()